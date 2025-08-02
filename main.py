import os
import asyncio
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from pathlib import Path

# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

# LangChain imports
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader, DirectoryLoader
from langchain_community.document_loaders.pdf import PyPDFLoader
from langchain_core.documents import Document

# Microsoft Word document processing
import docx2txt
from docx import Document as DocxDocument

# LangGraph imports
from langgraph.graph import StateGraph, START, END
from langgraph.prebuilt import ToolNode
from langchain_core.tools import tool

# Helper function to get file type information
def get_file_type_info(filename: str) -> str:
    """Get human-readable file type information"""
    ext = Path(filename).suffix.lower()
    file_types = {
        '.pdf': 'PDF Document',
        '.txt': 'Text File',
        '.doc': 'Microsoft Word Document',
        '.docx': 'Microsoft Word Document (Open XML)'
    }
    return file_types.get(ext, 'Unknown File Type')

# Custom Word Document Loader
class WordDocumentLoader:
    """Custom loader for Microsoft Word documents (.docx and .doc files)"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
    
    def load(self) -> List[Document]:
        """Load and extract text from Word documents"""
        try:
            print(f"Loading Word document: {self.file_path}")
            
            # Use docx2txt for better text extraction (handles both .doc and .docx)
            text = docx2txt.process(self.file_path)
            
            if not text or not text.strip():
                print(f"No text extracted with docx2txt, trying python-docx fallback...")
                # Fallback to python-docx for .docx files
                if self.file_path.lower().endswith('.docx'):
                    doc = DocxDocument(self.file_path)
                    text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            if not text or not text.strip():
                raise ValueError("No text content could be extracted from the document")
            
            print(f"Successfully extracted {len(text)} characters from Word document")
            
            # Create a Document object with metadata
            return [Document(
                page_content=text,
                metadata={
                    "source": self.file_path,
                    "file_type": "word_document",
                    "filename": Path(self.file_path).name
                }
            )]
            
        except Exception as e:
            print(f"Error loading Word document {self.file_path}: {str(e)}")
            raise Exception(f"Error loading Word document {self.file_path}: {str(e)}")

# FastAPI for web interface
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, validator
import uvicorn
import shutil

# Initialize FastAPI app
app = FastAPI(title="AI Chatbot with RAG", version="1.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# State management for LangGraph
@dataclass
class ChatState:
    documents: Optional[str] = None  # Retrieved relevant documents
    question: Optional[str] = None   # Current user question
    chat_history: List[Dict[str, str]] = None  # Conversation history
    response: Optional[str] = None   # Generated response
    error: Optional[str] = None      # Error message if any
    
    def __post_init__(self):
        if self.chat_history is None:
            self.chat_history = []
    
    def to_dict(self):
        """Convert ChatState to dictionary"""
        return {
            'documents': self.documents,
            'question': self.question,
            'chat_history': self.chat_history,
            'response': self.response,
            'error': self.error
        }

# Configuration
class Config:
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
    MODEL_NAME = "gemini-1.5-flash"
    CHROMA_PERSIST_DIRECTORY = "./chroma_db"
    DOCUMENTS_DIRECTORY = "./documents"
    MAX_TOKENS = 2048
    TEMPERATURE = 0.7

# Global components
llm = None
embeddings = None
vectorstore = None

# Initialize components
def initialize_components():
    """Initialize all components for the RAG system"""
    global llm, embeddings, vectorstore
    
    if llm is None:
        # Initialize Gemini LLM
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            google_api_key=Config.GEMINI_API_KEY,
            temperature=0,
        )
    
    if embeddings is None:
        # Initialize Google Gemini embeddings
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=Config.GEMINI_API_KEY
        )
    
    if vectorstore is None:
        # Initialize vector store
        vectorstore = Chroma(
            persist_directory=Config.CHROMA_PERSIST_DIRECTORY,
            embedding_function=embeddings
        )
    
    return llm, embeddings, vectorstore

# Document processing
def load_and_process_documents():
    """Load documents from directory and process them for vector store"""
    
    # Create documents directory if it doesn't exist
    Path(Config.DOCUMENTS_DIRECTORY).mkdir(exist_ok=True)
    
    documents = []
    
    # Load text documents
    try:
        text_loader = DirectoryLoader(
            Config.DOCUMENTS_DIRECTORY,
            glob="**/*.txt",
            loader_cls=TextLoader
        )
        documents.extend(text_loader.load())
    except Exception as e:
        print(f"No text documents found in {Config.DOCUMENTS_DIRECTORY}")
    
    # Load PDF documents
    try:
        pdf_loader = DirectoryLoader(
            Config.DOCUMENTS_DIRECTORY,
            glob="**/*.pdf",
            loader_cls=PyPDFLoader
        )
        documents.extend(pdf_loader.load())
    except Exception as e:
        print(f"No PDF documents found in {Config.DOCUMENTS_DIRECTORY}")
    
    # Load Word documents (.doc and .docx)
    try:
        word_files = list(Path(Config.DOCUMENTS_DIRECTORY).glob("**/*.doc*"))
        for word_file in word_files:
            try:
                word_loader = WordDocumentLoader(str(word_file))
                documents.extend(word_loader.load())
            except Exception as e:
                print(f"Error loading Word document {word_file}: {e}")
    except Exception as e:
        print(f"No Word documents found in {Config.DOCUMENTS_DIRECTORY}")            
    try:
        word_files = list(Path(Config.DOCUMENTS_DIRECTORY).glob("**/*.docx*"))
        for word_file in word_files:
            try:
                word_loader = WordDocumentLoader(str(word_file))
                documents.extend(word_loader.load())
            except Exception as e:
                print(f"Error loading Word document {word_file}: {e}")
    except Exception as e:
        print(f"No Word documents found in {Config.DOCUMENTS_DIRECTORY}")
    
    # If no documents found, create a sample document
    if not documents:
        print(f"No documents found in {Config.DOCUMENTS_DIRECTORY}. Creating sample document.")
        sample_doc_path = Path(Config.DOCUMENTS_DIRECTORY) / "sample.txt"
        with open(sample_doc_path, "w") as f:
            f.write("This is a sample document for the RAG system. It contains information about AI and machine learning.")
        
        # Load the sample document
        text_loader = DirectoryLoader(
            Config.DOCUMENTS_DIRECTORY,
            glob="**/*.txt",
            loader_cls=TextLoader
        )
        documents = text_loader.load()
    
    # Split documents
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    
    splits = text_splitter.split_documents(documents)
    return splits

# RAG Chain
def create_rag_chain(llm, vectorstore):
    """Create the RAG chain for question answering"""
    
    # Template for RAG
    template = """You are a helpful AI assistant. Use the following context to answer the user's question.
    If you don't know the answer, just say that you don't know, don't try to make up an answer.
    
    Context: {context}
    
    Question: {question}
    
    Answer:"""
    
    prompt = ChatPromptTemplate.from_template(template)
    
    # Create RAG chain
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)
    
    # Get retriever
    retriever = vectorstore.as_retriever(search_kwargs={"k": 3})
    
    rag_chain = (
        {"context": retriever | format_docs, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )
    
    return rag_chain

# LangGraph nodes
def retrieve_documents(state: ChatState) -> ChatState:
    """Retrieve relevant documents from vector store based on the question"""
    try:
        global llm, embeddings, vectorstore
        if vectorstore is None:
            initialize_components()
        
        # Retrieve relevant documents based on the question
        docs = vectorstore.similarity_search(state.question, k=3)
        documents = "\n\n".join([doc.page_content for doc in docs])
        
        # Create new state with updated documents
        new_state = ChatState(
            documents=documents,
            question=state.question,
            chat_history=state.chat_history or [],
            response=state.response,
            error=state.error
        )
        
        print(f"Debug: retrieve_documents returning: {type(new_state)}")
        return new_state
        
    except Exception as e:
        error_state = ChatState(
            documents=state.documents,
            question=state.question,
            chat_history=state.chat_history or [],
            response=state.response,
            error=str(e)
        )
        print(f"Debug: retrieve_documents error returning: {type(error_state)}")
        return error_state

def generate_response(state: ChatState) -> ChatState:
    """Generate response using RAG chain with documents, question, and chat history"""
    try:
        global llm, embeddings, vectorstore
        if llm is None or vectorstore is None:
            initialize_components()
        
        # Create enhanced prompt with chat history
        template = """You are a helpful AI assistant. Use the following context to answer the user's question.
        Consider the conversation history to provide contextual and relevant responses.
        
        Conversation History:
        {chat_history}
        
        Relevant Documents:
        {documents}
        
        Current Question: {question}
        
        Answer:"""
        
        prompt = ChatPromptTemplate.from_template(template)
        
        # Format chat history
        chat_history_text = ""
        if state.chat_history:
            for msg in state.chat_history:
                sender = msg.get("sender", "user")
                content = msg.get("content", "")
                # Map "ai" sender to "Assistant" for better readability
                display_sender = "Assistant" if sender == "ai" else sender.capitalize()
                chat_history_text += f"{display_sender}: {content}\n"
        
        # Generate response
        response = llm.invoke(prompt.format(
            chat_history=chat_history_text,
            documents=state.documents or "No relevant documents found.",
            question=state.question
        ))
        
        # Create new state with response
        new_state = ChatState(
            documents=state.documents,
            question=state.question,
            chat_history=state.chat_history or [],
            response=response.content,
            error=state.error
        )
        
        print(f"Debug: generate_response returning: {type(new_state)}")
        return new_state
        
    except Exception as e:
        error_state = ChatState(
            documents=state.documents,
            question=state.question,
            chat_history=state.chat_history or [],
            response=state.response,
            error=str(e)
        )
        print(f"Debug: generate_response error returning: {type(error_state)}")
        return error_state

# Create LangGraph workflow
def create_chatbot_workflow():
    """Create the LangGraph workflow for the chatbot"""
    
    # Create the graph
    workflow = StateGraph(ChatState)
    
    # Add nodes
    workflow.add_node("retrieve_documents", retrieve_documents)
    workflow.add_node("generate_response", generate_response)
    
    # Add edges
    workflow.add_edge(START, "retrieve_documents")
    workflow.add_edge("retrieve_documents", "generate_response")
    workflow.add_edge("generate_response", END)
    
    # Compile the graph
    app = workflow.compile()
    
    return app

# Initialize the workflow
chatbot_workflow = create_chatbot_workflow()

# Pydantic models for API
class ChatMessage(BaseModel):
    sender: str  # "user" or "ai"
    content: str
    
    @validator('sender')
    def validate_sender(cls, v):
        if v not in ['user', 'ai']:
            raise ValueError('sender must be either "user" or "ai"')
        return v

class ChatRequest(BaseModel):
    question: str
    chat_history: Optional[List[ChatMessage]] = []
    
    @validator('question')
    def validate_question(cls, v):
        if not v or not v.strip():
            raise ValueError('question cannot be empty')
        return v.strip()

class ChatResponse(BaseModel):
    response: str
    documents: Optional[str] = None
    error: Optional[str] = None
    
    class Config:
        from_attributes = True

# API endpoints
@app.post("/chat")
async def chat(request: ChatRequest):
    """Chat endpoint for the RAG chatbot"""
    print(f"------>>>>>>>>>Question is {request.question.strip()}")
    try:
        print(f"------>>>>>>>>>Question is {request.question.strip()}")
        global system_initialized, llm, embeddings, vectorstore
        
        # Initialize system if not already done
        if not system_initialized:
            system_initialized = initialize_system()
            if not system_initialized:
                raise HTTPException(
                    status_code=500, 
                    detail="RAG system not initialized. Please set a valid GEMINI_API_KEY environment variable."
                )
        
        # Convert ChatMessage objects to dict format for internal use
        chat_history_dict = []
        if request.chat_history:
            for msg in request.chat_history:
                chat_history_dict.append({
                    "sender": msg.sender,
                    "content": msg.content
                })
        
        # Create initial state with the three required states
        initial_state = ChatState(
            documents=None,  # Will be populated by retrieve_documents node
            question=request.question,
            chat_history=chat_history_dict,
            response=None,
            error=None
        )
        
        # Run the workflow
        result = chatbot_workflow.invoke(initial_state)
        
        # Debug: Print the result type and content
        print(f"Debug: Result type: {type(result)}")
        print(f"Debug: Result content: {result}")
        
        # Handle both dict and ChatState results
        if isinstance(result, dict):
            response = result.get('response', "I'm sorry, I couldn't generate a response.")
            documents = result.get('documents')
            error = result.get('error')
        else:
            response = result.response or "I'm sorry, I couldn't generate a response."
            documents = result.documents
            error = result.error
        
        # Return JSON response
        return JSONResponse(
            status_code=200,
            content={
                "answer": response,
                "documents": documents,
                "error": error,
                "status": "success" if not error else "error"
            }
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-file")
async def upload_file(file: UploadFile = File(...)):
    """Upload a single file (PDF or TXT) to the knowledge base"""
    try:
        global vectorstore
        if vectorstore is None:
            initialize_components()
        
        # Validate that file was actually received
        if not file or not file.filename:
            raise HTTPException(status_code=400, detail="No file provided or invalid file")
        
        # Validate file type
        if not file.filename.lower().endswith(('.pdf', '.txt', '.doc', '.docx')):
            raise HTTPException(status_code=400, detail="Only PDF, TXT, DOC, and DOCX files are supported")
        
        # Validate file size (max 10MB)
        if file.size and file.size > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="File size must be less than 10MB")
        
        # Create documents directory if it doesn't exist
        Path(Config.DOCUMENTS_DIRECTORY).mkdir(exist_ok=True)
        
        # Generate unique filename to avoid conflicts
        import uuid
        file_extension = Path(file.filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = Path(Config.DOCUMENTS_DIRECTORY) / unique_filename
        
        # Save uploaded file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Load and process the uploaded file
        try:
            if file.filename.lower().endswith('.pdf'):
                loader = PyPDFLoader(str(file_path))
            elif file.filename.lower().endswith(('.doc', '.docx')):
                loader = WordDocumentLoader(str(file_path))
            else:
                loader = TextLoader(str(file_path))
            
            documents = loader.load()
            
            if not documents:
                raise HTTPException(status_code=400, detail="No content could be extracted from the file")
            
            # Split documents
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
            )
            
            splits = text_splitter.split_documents(documents)
            
            # Add to vector store
            vectorstore.add_documents(splits)
            
            return JSONResponse(
                status_code=200,
                content={
                    "message": f"Successfully uploaded and processed {file.filename}",
                    "filename": unique_filename,
                    "chunks": len(splits),
                    "original_filename": file.filename,
                    "file_type": get_file_type_info(file.filename),
                    "file_size": file.size
                }
            )
            
        except Exception as processing_error:
            # Clean up the uploaded file if processing fails
            if file_path.exists():
                file_path.unlink()
            
            # Provide more specific error messages for different file types
            error_msg = str(processing_error)
            if "word document" in error_msg.lower():
                raise HTTPException(
                    status_code=500, 
                    detail=f"Error processing Word document: {error_msg}. Please ensure the file is not corrupted and contains readable text."
                )
            else:
                raise HTTPException(status_code=500, detail=f"Error processing file: {error_msg}")
    
    except HTTPException:
        raise
    except Exception as e:
        error_msg = str(e)
        if "boundary" in error_msg.lower():
            raise HTTPException(
                status_code=400, 
                detail="Invalid request format. Please ensure you're sending a proper multipart/form-data request with a file."
            )
        else:
            raise HTTPException(status_code=500, detail=f"Upload failed: {error_msg}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    global system_initialized
    
    # Check if API key is set
    api_key = os.getenv("GEMINI_API_KEY")
    api_key_status = "set" if api_key and api_key != "your-gemini-api-key-here" else "not set"
    
    return JSONResponse(
        status_code=200,
        content={
            "status": "healthy" if system_initialized else "initializing",
            "message": "AI Chatbot is running",
            "api_key": api_key_status,
            "rag_system": "initialized" if system_initialized else "not initialized"
        }
    )

@app.get("/files")
async def list_files():
    """List all uploaded files in the documents directory"""
    try:
        documents_path = Path(Config.DOCUMENTS_DIRECTORY)
        if not documents_path.exists():
            return JSONResponse(status_code=200, content={"files": []})
        
        files = []
        for file_path in documents_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in ['.pdf', '.txt', '.doc', '.docx']:
                files.append({
                    "filename": file_path.name,
                    "size": file_path.stat().st_size,
                    "uploaded": file_path.stat().st_mtime,
                    "file_type": get_file_type_info(file_path.name)
                })
        
        return JSONResponse(status_code=200, content={"files": files})
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/files/{filename}")
async def delete_file(filename: str):
    """Delete a specific file from the documents directory"""
    try:
        file_path = Path(Config.DOCUMENTS_DIRECTORY) / filename
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found")
        
        if file_path.suffix.lower() not in ['.pdf', '.txt', '.doc', '.docx']:
            raise HTTPException(status_code=400, detail="Invalid file type")
        
        # Remove the file
        file_path.unlink()
        
        return JSONResponse(
            status_code=200,
            content={"message": f"Successfully deleted {filename}"}
        )
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Tools for LangGraph (optional)
@tool
def search_documents(query: str) -> str:
    """Search for relevant documents in the knowledge base"""
    try:
        global vectorstore
        if vectorstore is None:
            initialize_components()
        
        docs = vectorstore.similarity_search(query, k=3)
        return "\n\n".join([doc.page_content for doc in docs])
    except Exception as e:
        return f"Error searching documents: {str(e)}"

# Helper function to format chat history for display
def format_chat_history(chat_history: List[Dict[str, str]]) -> str:
    """Format chat history for display"""
    if not chat_history:
        return "No previous conversation."
    
    formatted = []
    for msg in chat_history:
        sender = msg.get("sender", "user")
        content = msg.get("content", "")
        # Map "ai" sender to "Assistant" for better readability
        display_sender = "Assistant" if sender == "ai" else sender.capitalize()
        formatted.append(f"{display_sender}: {content}")
    
    return "\n".join(formatted)

# Initialize the system
def initialize_system():
    """Initialize the RAG system components and load documents"""
    print("Initializing AI Chatbot with RAG system...")
    
    try:
        # Check if API key is valid first
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key or api_key == "your_actual_api_key_here":
            print("❌ GEMINI_API_KEY not set or using default value")
            print("Please set your Gemini API key in the .env file or environment variable")
            return False
        
        # Initialize components
        initialize_components()
        
        # Test the embeddings with a simple call
        try:
            test_embedding = embeddings.embed_query("test")
            print("✅ API key validation successful")
        except Exception as api_error:
            print(f"❌ API key validation failed: {api_error}")
            return False
        
        # Load and process documents
        print("Loading and processing documents...")
        splits = load_and_process_documents()
        
        # Add documents to vector store
        vectorstore.add_documents(splits)
               
        print(f"Processed {len(splits)} document chunks")
        print("RAG system initialized successfully!")
        return True
    except Exception as e:
        print(f"Warning: Could not initialize RAG system: {e}")
        print("Please check your GEMINI_API_KEY and try again")
        return False

# Don't initialize at import time - let it happen when needed
system_initialized = False

# For direct execution with uvicorn
if __name__ == "__main__":
    print("Starting FastAPI server...")
    print("You can also run with: uvicorn main:app --host 0.0.0.0 --port 8000 --reload")
    uvicorn.run(app, host="0.0.0.0", port=8000)
